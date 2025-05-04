from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsmap
import pandas as pd

# Criar DataFrame de exemplo
df = pd.DataFrame({
    "Título": ["Lote 1", "Lote 2", "Lote 3"],
    "Descrição": [
        "Descrição completa do primeiro lote.",
        "Este lote contém itens exclusivos.",
        "Lote especial com condições diferenciadas."
    ]
})

# Criar documento Word
doc = Document()

# Definir margens
sections = doc.sections
for section in sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

# Percorrer DataFrame e adicionar títulos e descrições formatadas
for index, row in df.iterrows():
    # Título
    titulo_paragrafo = doc.add_paragraph(row["Título"])
    titulo_paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    titulo_paragrafo.paragraph_format.space_after = Pt(0)
    titulo_run = titulo_paragrafo.runs[0]
    titulo_run.font.name = "Tenorite"
    titulo_run.font.size = Pt(10)

    # Parágrafo em branco
    doc.add_paragraph("")

    # Descrição com recuo completo e barra vertical
    descricao_paragrafo = doc.add_paragraph(row["Descrição"])
    descricao_paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    descricao_paragrafo.paragraph_format.left_indent = Cm(1)  # Recuo total do parágrafo
    descricao_paragrafo.paragraph_format.space_after = Pt(0)

    # Adicionar borda vertical à esquerda (barra)
    p_xml = descricao_paragrafo._element
    pPr = p_xml.get_or_add_pPr()
    pBdr = parse_xml('<w:pBdr xmlns:w="{}"><w:left w:val="single" w:sz="8" w:space="4" w:color="000000"/></w:pBdr>'.format(nsmap['w']))
    pPr.append(pBdr)

    descricao_run = descricao_paragrafo.runs[0]
    descricao_run.font.name = "Tenorite"
    descricao_run.font.size = Pt(10)

    # Dois parágrafos em branco antes do próximo título
    doc.add_paragraph("")
    doc.add_paragraph("")

# Salvar documento Word
doc.save("C:\\Users\\Yurih Santos\\Documents\\dados_formatado.docx")